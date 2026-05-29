from __future__ import annotations

from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .models import FormatStats, ParagraphInfo, ParagraphRole, ParagraphStyleRule, TableRule, TemplateRules


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

DEFAULT_WORD_STYLE_NAMES = {
    ParagraphRole.TITLE: "论文题目",
    ParagraphRole.ABSTRACT_HEADING: "摘要标题",
    ParagraphRole.ABSTRACT: "摘要正文",
    ParagraphRole.KEYWORDS: "关键词",
    ParagraphRole.HEADING_1: "一级标题",
    ParagraphRole.HEADING_2: "二级标题",
    ParagraphRole.HEADING_3: "三级标题",
    ParagraphRole.TOC_ENTRY: "目录条目",
    ParagraphRole.FIGURE_CAPTION: "图题",
    ParagraphRole.TABLE_CAPTION: "表题",
    ParagraphRole.REFERENCES_HEADING: "参考文献标题",
    ParagraphRole.REFERENCE: "参考文献条目",
    ParagraphRole.BODY: "正文",
}


def apply_template(document: DocumentObject, structure: list[ParagraphInfo], rules: TemplateRules) -> FormatStats:
    stats = FormatStats(total_paragraphs=len(structure))
    apply_page_rules(document, rules)
    word_styles = ensure_template_styles(document, rules)

    for info in structure:
        stats.role_counts[info.role] = stats.role_counts.get(info.role, 0) + 1
        rule = rules.styles.get(info.role) or rules.styles.get(ParagraphRole.BODY)
        if info.role == ParagraphRole.EMPTY or rule is None:
            continue
        style_name = word_styles.get(info.role) or word_styles.get(ParagraphRole.BODY)
        apply_paragraph_style(document.paragraphs[info.index], style_name)
        stats.styled_paragraphs += 1

    add_basic_warnings(stats)
    stats.formatted_tables = apply_table_rules(document, rules)
    return stats


def apply_page_rules(document: DocumentObject, rules: TemplateRules) -> None:
    for section in document.sections:
        if rules.page.margin_top_cm is not None:
            section.top_margin = Cm(rules.page.margin_top_cm)
        if rules.page.margin_bottom_cm is not None:
            section.bottom_margin = Cm(rules.page.margin_bottom_cm)
        if rules.page.margin_left_cm is not None:
            section.left_margin = Cm(rules.page.margin_left_cm)
        if rules.page.margin_right_cm is not None:
            section.right_margin = Cm(rules.page.margin_right_cm)


def ensure_template_styles(
    document: DocumentObject, rules: TemplateRules
) -> dict[ParagraphRole, str]:
    style_names: dict[ParagraphRole, str] = {}
    for role, rule in rules.styles.items():
        if role == ParagraphRole.EMPTY:
            continue
        style_name = rule.word_style_name or DEFAULT_WORD_STYLE_NAMES.get(role, role.value)
        style = get_or_create_paragraph_style(document, style_name)
        apply_rule_to_word_style(style, rule)
        style_names[role] = style_name
    return style_names


def get_or_create_paragraph_style(document: DocumentObject, style_name: str):
    try:
        return document.styles[style_name]
    except KeyError:
        return document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def apply_rule_to_word_style(style, rule: ParagraphStyleRule) -> None:
    fmt = style.paragraph_format
    if rule.alignment:
        fmt.alignment = ALIGNMENTS.get(rule.alignment)
    if rule.first_line_indent_cm is not None:
        fmt.first_line_indent = Cm(rule.first_line_indent_cm)
    if rule.left_indent_cm is not None:
        fmt.left_indent = Cm(rule.left_indent_cm)
    if rule.line_spacing_pt is not None:
        fmt.line_spacing = Pt(rule.line_spacing_pt)
    if rule.line_spacing is not None:
        fmt.line_spacing = rule.line_spacing
    if rule.space_before_pt is not None:
        fmt.space_before = Pt(rule.space_before_pt)
    if rule.space_after_pt is not None:
        fmt.space_after = Pt(rule.space_after_pt)
    if rule.page_break_before is not None:
        fmt.page_break_before = rule.page_break_before
    if rule.keep_with_next is not None:
        fmt.keep_with_next = rule.keep_with_next
    if rule.keep_together is not None:
        fmt.keep_together = rule.keep_together

    if rule.font:
        style.font.name = rule.font
    if rule.east_asia_font:
        set_style_east_asia_font(style, rule.east_asia_font)
    if rule.size_pt is not None:
        style.font.size = Pt(rule.size_pt)
    if rule.bold is not None:
        style.font.bold = rule.bold
    if rule.italic is not None:
        style.font.italic = rule.italic
    style.quick_style = True
    style.hidden = False
    style.unhide_when_used = True
    if style.priority is None:
        style.priority = 10


def set_style_east_asia_font(style, font_name: str) -> None:
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def apply_paragraph_style(paragraph, style_name: str | None) -> None:
    if style_name is None:
        return
    paragraph.style = style_name


def apply_table_rules(document: DocumentObject, rules: TemplateRules) -> int:
    if rules.table is None:
        return 0

    table_style = get_or_create_paragraph_style(document, "表内文字")
    apply_table_text_rule_to_style(table_style, rules.table)

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        apply_three_line_borders(table, rules.table)
        for row in table.rows:
            if rules.table.row_height_cm is not None:
                row.height = Cm(rules.table.row_height_cm)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.style = table_style
    return len(document.tables)


def apply_table_text_rule_to_style(style, rule: TableRule) -> None:
    fmt = style.paragraph_format
    if rule.alignment:
        fmt.alignment = ALIGNMENTS.get(rule.alignment)
    if rule.line_spacing is not None:
        fmt.line_spacing = rule.line_spacing
    if rule.font:
        style.font.name = rule.font
    if rule.east_asia_font:
        set_style_east_asia_font(style, rule.east_asia_font)
    if rule.size_pt is not None:
        style.font.size = Pt(rule.size_pt)
    style.quick_style = True
    style.hidden = False
    style.unhide_when_used = True
    if style.priority is None:
        style.priority = 10


def apply_three_line_borders(table, rule: TableRule) -> None:
    if not table.rows:
        return
    clear_all_table_cell_borders(table)
    set_row_border(table.rows[0], "top", rule.top_border_pt or 1.5)
    set_row_border(table.rows[0], "bottom", rule.header_bottom_border_pt or 0.75)
    set_row_border(table.rows[-1], "bottom", rule.bottom_border_pt or 1.5)


def clear_all_table_cell_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = get_or_add_child(tc_pr, "w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = get_or_add_child(tc_borders, f"w:{edge}")
                border.set(qn("w:val"), "nil")
                border.set(qn("w:sz"), "0")


def set_row_border(row, edge: str, width_pt: float) -> None:
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = get_or_add_child(tc_pr, "w:tcBorders")
        border = get_or_add_child(tc_borders, f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(int(width_pt * 8)))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def add_basic_warnings(stats: FormatStats) -> None:
    if stats.role_counts.get(ParagraphRole.ABSTRACT, 0) == 0:
        stats.warnings.append("No abstract paragraph was detected.")
    if stats.role_counts.get(ParagraphRole.KEYWORDS, 0) == 0:
        stats.warnings.append("No keywords paragraph was detected.")
    if stats.role_counts.get(ParagraphRole.REFERENCES_HEADING, 0) == 0:
        stats.warnings.append("No references heading was detected.")
