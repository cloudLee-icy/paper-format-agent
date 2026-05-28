from pathlib import Path

from docx import Document


def main() -> None:
    output = Path(__file__).with_name("messy_paper.docx")
    doc = Document()
    doc.add_paragraph("面向学术写作的文档格式化方法研究")
    doc.add_paragraph("摘要：本文讨论一种将不规范论文文档转换为目标模板格式的方法。")
    doc.add_paragraph("关键词：论文格式化；Word；模板；学术写作")
    doc.add_paragraph("1 引言")
    doc.add_paragraph("作者在投稿或提交毕业论文前，通常需要花费大量时间处理格式问题。")
    doc.add_paragraph("1.1 研究背景")
    doc.add_paragraph("不同期刊和学校对字体、行距、页边距、标题层级和参考文献有不同要求。")
    doc.add_paragraph("图 1 文档格式化流程")
    doc.add_paragraph("2 方法")
    doc.add_paragraph("本文采用结构识别和确定性格式套用相结合的方式。")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Zhang S. Academic document formatting workflow. 2026.")
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()

