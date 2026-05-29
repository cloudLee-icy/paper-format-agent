from pathlib import Path

from docx import Document


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def main() -> None:
    output = Path(__file__).with_name("full_thesis_demo.docx")
    doc = Document()

    add_paragraphs(
        doc,
        [
            "面向学术论文格式化的模板驱动方法研究",
            "作者姓名：张三",
            "学院：电子科学与工程学院",
            "专业：电子信息",
            "学号：2024000000",
            "指导教师：李四 教授",
            "",
            "独创性声明",
            "本人声明所呈交的学位论文为本人在导师指导下独立完成的研究成果，文中引用内容均已注明来源。",
            "论文使用授权",
            "本人同意学校按照相关规定保存、使用和提供本学位论文。",
            "",
            "摘要",
            "摘要：学术论文在投稿和毕业提交前通常需要满足严格的格式要求。本文提出一种模板驱动的文档格式化方法，将论文结构识别、Word 样式生成和格式检查报告结合起来，以提高格式整理的可复现性和可审计性。",
            "关键词：学位论文；文档格式化；Word 样式；模板驱动；格式检查",
            "",
            "ABSTRACT",
            "Abstract: Academic manuscripts usually need to follow strict formatting requirements before submission. This thesis proposes a template-driven formatting workflow that combines structure detection, Word style generation, and auditable reports.",
            "Keywords: thesis; document formatting; Word styles; template-driven workflow; format report",
            "",
            "目录",
            "第一章 绪论",
            "第二章 相关技术",
            "第三章 模板驱动格式化方法",
            "第四章 实验与分析",
            "结论",
            "致谢",
            "参考文献",
            "附录A 示例模板字段",
            "",
            "图目录",
            "图 1 文档格式化流程",
            "表目录",
            "表 1 模板字段示例",
            "主要符号表",
            "T 表示目标格式模板",
            "D 表示待格式化文档",
            "缩略词表",
            "DOCX Microsoft Word Open XML 文档格式",
            "",
            "第一章 绪论",
            "1.1 研究背景",
            "论文作者在整理格式时经常需要反复调整页边距、标题、正文、图表题注和参考文献格式。人工处理过程耗时，且不同人员之间的结果不容易保持一致。",
            "1.2 研究意义",
            "将格式要求转换为可执行模板，并通过程序生成 Word 样式，可以降低重复劳动，同时保留人工继续修改样式的空间。",
            "1.3 本文主要工作",
            "本文围绕论文结构识别、模板规则表达、Word 样式应用和检查报告生成四个方面展开。",
            "",
            "第二章 相关技术",
            "2.1 Word 样式机制",
            "Word 样式可以集中控制同类段落的字体、字号、缩进、行距和段前段后距，是论文格式化中最适合复用的基础机制。",
            "2.2 模板规则表达",
            "模板规则可以来自学校规范、期刊投稿指南、Word 模板、Markdown 模板或 LaTeX/Overleaf 模板。",
            "",
            "第三章 模板驱动格式化方法",
            "3.1 方法流程",
            "本文方法首先读取待处理文档，再识别段落角色，随后根据模板创建对应 Word 样式，最后将段落应用到相应样式并输出检查报告。",
            "图 1 文档格式化流程",
            "3.2 模板字段设计",
            "模板字段需要覆盖页面设置、标题层级、正文、摘要、关键词、图题、表题和参考文献条目等内容。",
            "表 1 模板字段示例",
        ],
    )

    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "含义"
    table.cell(0, 2).text = "示例"
    table.cell(1, 0).text = "heading_1"
    table.cell(1, 1).text = "一级标题"
    table.cell(1, 2).text = "第一章 绪论"
    table.cell(2, 0).text = "body"
    table.cell(2, 1).text = "正文"
    table.cell(2, 2).text = "宋体小四"

    add_paragraphs(
        doc,
        [
            "3.3 公式示例",
            "F = m × a                                                (3-1)",
            "式中，F 表示力，m 表示质量，a 表示加速度。",
            "",
            "第四章 实验与分析",
            "4.1 实验设计",
            "本文构造包含封面信息、声明、摘要、目录、正文、图表、公式、参考文献、附录和成果的完整示例论文，用于检验格式化流程。",
            "4.2 结果分析",
            "实验关注段落角色是否被正确识别，以及输出文档中是否生成内容匹配的 Word 样式。",
            "",
            "结论",
            "本文构建了一个模板驱动的论文格式化示例流程。实验表明，基于内容角色创建 Word 样式有利于自动格式化和后续人工调整。",
            "",
            "致谢",
            "感谢导师和同学在论文写作与工具验证过程中提供的帮助。",
            "",
            "参考文献",
            "[1] 全国信息与文献标准化技术委员会. 信息与文献 参考文献著录规则: GB/T 7714-2015[S]. 北京: 中国标准出版社, 2015.",
            "[2] Zhang S. Template-driven academic document formatting workflow[J]. Journal of Scholarly Tools, 2026, 1(1): 1-12.",
            "[3] Microsoft. Word styles and document formatting guide[EB/OL]. 2026.",
            "",
            "附录A 示例模板字段",
            "本附录列出模板中常见字段，包括正文、一级标题、二级标题、图题、表题和参考文献条目。",
            "",
            "攻读学位期间取得的成果",
            "[1] 张三. 面向学术写作的文档格式化工具设计[C]. 学术工具与出版技术会议, 2026.",
        ],
    )

    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
